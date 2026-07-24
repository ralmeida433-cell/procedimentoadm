"""Motor de Análise e Extração de Documentos (SAD/PCD/SI/RIP/Ata do CEDMU).

Recebe um arquivo (PDF, DOCX ou imagem), extrai o texto (com OCR via
Tesseract quando o PDF não tem texto embutido, ex.: digitalização) e usa a
API do OpenRouter para classificar o tipo de procedimento e extrair as
entidades estruturadas no formato do schema canônico do sistema
(`pcd_automation.schema`).

Mesmo princípio do resto do sistema: a IA nunca inventa um dado que não
esteja no texto - campos não encontrados ficam de fora do JSON (o
chamador decide o que fazer com a ausência), e ambiguidades são reportadas
explicitamente em vez de "resolvidas" silenciosamente.
"""
from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field
from pathlib import Path

from .ia_cliente import chamar_openrouter

EXTENSOES_SUPORTADAS = {".pdf", ".docx", ".png", ".jpg", ".jpeg"}

# Campos do schema canônico (ver pcd_automation.schema) que valem a pena
# tentar extrair de um documento já redigido - nem todo documento terá
# todos eles, e a IA é instruída a só preencher o que encontrar.
CAMPOS_ALVO = [
    "reds", "nome_sindicado", "re_sindicado", "posto_graduacao_sindicado", "unidade_sindicado",
    "data_fato", "cidade_fato", "hora_fato", "local_fato", "resumo_fato",
    "nome_autoridade_processante", "posto_autoridade_processante", "re_autoridade_processante",
    "unidade_autoridade_processante",
    "nome_autoridade_delegante", "posto_autoridade_delegante",
    "data_instauracao", "numero_processo", "numero_regiao_pm", "numero_batalhao_pm", "cidade_sede",
    "nome_comunicante", "posto_comunicante", "re_comunicante", "unidade_comunicante", "data_comunicacao",
    "numero_comunicacao_disciplinar", "numero_folhas_comunicacao_disciplinar", "numero_folhas_escala_servico",
    "tipificacao_cedm",
    "nome_testemunha", "posto_testemunha", "re_testemunha", "unidade_testemunha",
]

PROMPT_SISTEMA = f"""Você é o Motor de Análise e Extração de Documentos do Especialista MAPPA (PMMG/CBMMG).

FASE 1 - Classifique o tipo de procedimento do texto recebido: SAD (Sindicância Administrativa \
Disciplinar), PCD (Processo/Processo de Comunicação Disciplinar), SI (Sindicância Investigatória), \
RIP (Relatório de Informação Policial), "Ata de Reunião do CEDMU", ou "outro" se não se encaixar em \
nenhum desses. Se a classificação não for clara, diga isso no campo "confianca_tipo": "baixa" em vez \
de arriscar.

FASE 2 - Extraia os dados estruturados encontrados no texto para os campos a seguir (schema canônico \
do sistema). NUNCA invente um valor que não esteja no texto - se não encontrar um campo, OMITA-O do \
JSON (não escreva null nem string vazia). Campos-alvo: {", ".join(CAMPOS_ALVO)}.

Formatação obrigatória dos valores (convenções específicas deste sistema, siga rigorosamente):
- Datas: formato ISO "aaaa-mm-dd".
- "hora_fato": só o número da hora, sem "h"/"min"/"horas" (ex.: "23", nunca "23h30min" nem "23:30").
- "numero_processo": só o número sequencial, SEM o ano (ex.: "010", nunca "010/2026" - o ano é \
derivado de "data_instauracao").
- "resumo_fato": no particípio, como continuação de "teria o comunicado ___" (ex.: "deixado de \
prestar continência...", "exercido atividade de segurança privada...").
- "numero_regiao_pm" e "numero_batalhao_pm": só o número, sem "ª"/"º" nem "Região"/"Batalhão" (ex.: \
"29" e "101", nunca "29ª Região" nem "101º BPM") - normalmente aparecem no cabeçalho/timbre de cada \
página do documento ("29ª REGIÃO DE POLÍCIA MILITAR", "101º BATALHÃO DE POLÍCIA MILITAR").

Se houver ambiguidade (nomes/unidades divergentes entre partes do texto, trecho ilegível ou cortado, \
mais de uma pessoa possível para o mesmo papel), NÃO tente resolver sozinho - liste o problema em \
"ambiguidades".

Responda SOMENTE com um objeto JSON válido, sem markdown, sem texto antes ou depois, exatamente neste \
formato:
{{"tipo_documento": "...", "confianca_tipo": "alta|media|baixa", "campos": {{"...": "..."}}, "ambiguidades": ["..."]}}
"""


@dataclass
class ResultadoExtracao:
    tipo_documento: str | None = None
    confianca_tipo: str | None = None
    campos: dict = field(default_factory=dict)
    ambiguidades: list[str] = field(default_factory=list)
    texto_extraido: str = ""
    erro: str | None = None


def extrair_texto(caminho: Path) -> str:
    """Extrai o texto de um PDF, DOCX ou imagem. Levanta ValueError para
    extensões não suportadas."""
    sufixo = caminho.suffix.lower()
    if sufixo == ".pdf":
        return _extrair_texto_pdf(caminho)
    if sufixo == ".docx":
        return _extrair_texto_docx(caminho)
    if sufixo in {".png", ".jpg", ".jpeg"}:
        return _extrair_texto_imagem(caminho)
    raise ValueError(
        f"Extensão não suportada: {sufixo} (aceitos: {', '.join(sorted(EXTENSOES_SUPORTADAS))})"
    )


def _extrair_texto_pdf(caminho: Path) -> str:
    import fitz  # PyMuPDF

    partes = []
    with fitz.open(str(caminho)) as doc:
        for pagina in doc:
            texto_pagina = pagina.get_text().strip()
            if not texto_pagina:
                # Página sem texto embutido (provável digitalização/imagem) - cai para OCR.
                texto_pagina = _ocr_pixmap(pagina.get_pixmap(dpi=200))
            partes.append(texto_pagina)
    return "\n\n".join(partes)


def _extrair_texto_docx(caminho: Path) -> str:
    import docx

    d = docx.Document(str(caminho))
    partes = [p.text for p in d.paragraphs if p.text.strip()]
    for tabela in d.tables:
        for linha in tabela.rows:
            texto_linha = " | ".join(c.text.strip() for c in linha.cells)
            if texto_linha.strip(" |"):
                partes.append(texto_linha)
    return "\n".join(partes)


def _extrair_texto_imagem(caminho: Path) -> str:
    import pytesseract
    from PIL import Image

    with Image.open(caminho) as imagem:
        return pytesseract.image_to_string(imagem, lang="por")


def _ocr_pixmap(pixmap) -> str:
    import pytesseract
    from PIL import Image

    with Image.open(io.BytesIO(pixmap.tobytes("png"))) as imagem:
        return pytesseract.image_to_string(imagem, lang="por")


def _extrair_json(texto: str) -> dict:
    texto = texto.strip()
    # Remove cercas de markdown (```json ... ```) caso a IA ignore a instrução de "só JSON".
    texto = re.sub(r"^```(?:json)?\s*|\s*```$", "", texto, flags=re.IGNORECASE)
    return json.loads(texto)


def classificar_e_extrair(texto: str) -> ResultadoExtracao:
    resultado = ResultadoExtracao(texto_extraido=texto)
    if not texto.strip():
        resultado.erro = (
            "Não foi possível extrair nenhum texto do arquivo (pode estar em branco, corrompido ou "
            "o OCR não reconheceu nada)."
        )
        return resultado

    conteudo, erro = chamar_openrouter(
        [
            {"role": "system", "content": PROMPT_SISTEMA},
            # O modelo padrão tem 1M de tokens de contexto - 150k caracteres cobre
            # com folga até processos bem mais longos que um PCD/SAD típico.
            {"role": "user", "content": f"Texto extraído do documento:\n\n{texto[:150000]}"},
        ],
        max_tokens=6000,
        timeout=120,
        json_obrigatorio=True,
    )
    if erro:
        resultado.erro = erro
        return resultado

    try:
        dados_json = _extrair_json(conteudo)
    except json.JSONDecodeError:
        resultado.erro = f"A IA não retornou um JSON válido. Resposta bruta: {conteudo[:500]}"
        return resultado

    resultado.tipo_documento = dados_json.get("tipo_documento")
    resultado.confianca_tipo = dados_json.get("confianca_tipo")
    resultado.campos = dados_json.get("campos") or {}
    resultado.ambiguidades = dados_json.get("ambiguidades") or []
    return resultado
