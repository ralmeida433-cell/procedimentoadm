"""Constrói o template docxtpl da Proposta de Recompensa a partir do modelo
oficial da PMMG.

Fonte: `Modelo Processo-Procedimento/PROPOSTA DE RECOMPENSA/*.doc` (modelo
oficial, NUNCA alterado). Como é um .doc binário, a conversão para .docx é
feita fora daqui uma única vez (nesta máquina: o conversor do Office 2007,
`Office12/Wordconv.exe -oice -nme entrada.doc saida.docx`) e este script
transforma a cópia convertida no template `templates/proposta_recompensa.docx`.

O modelo é um exemplo PREENCHIDO (REDS 2025-013916881-001, 15 militares).
Este script preserva a formatação parágrafo a parágrafo e:

1. troca os dados do caso de exemplo por marcadores Jinja ({{ ... }});
2. reduz os 15 blocos de individualização/critérios a UM bloco dentro de um
   laço {%p for m in militares %} (o docxtpl repete o bloco formatado para
   cada militar na geração);
3. remove os anexos do caso (fotos/reportagens de Nova Lima) e deixa um laço
   de linhas de link no lugar;
4. remove do pacote as imagens que ficaram órfãs (as fotos do caso), mantendo
   o brasão - o template cai de ~6 MB para poucos KB.

Rodar de novo só é preciso se o modelo oficial mudar:
    python -m pcd_automation.modelo_documentos.gerar_template_recompensa <convertido.docx>
"""
from __future__ import annotations

import copy
import re
import shutil
import sys
import zipfile
from pathlib import Path

import docx
from docx.oxml.ns import qn

DESTINO = Path(__file__).resolve().parent / "templates" / "proposta_recompensa.docx"


def _definir_texto(paragrafo, texto: str) -> None:
    """Troca o texto do parágrafo preservando a formatação do primeiro run."""
    runs = paragrafo.runs
    if not runs:
        paragrafo.add_run(texto)
        return
    runs[0].text = texto
    for run in runs[1:]:
        run.text = ""


def _paragrafo_tag(modelo_par, tag: str):
    """Cria um parágrafo novo (cópia do modelo) contendo só a tag Jinja."""
    novo = copy.deepcopy(modelo_par._p)
    par = docx.text.paragraph.Paragraph(novo, modelo_par._parent)
    _definir_texto(par, tag)
    return novo


def _texto_de(el) -> str:
    return "".join(n.text or "" for n in el.iter(qn("w:t"))).strip()


def construir(origem: Path, destino: Path = DESTINO) -> None:
    d = docx.Document(str(origem))
    corpo = d.element.body
    filhos = list(corpo)

    def indice_par(prefixo: str, a_partir: int = 0) -> int:
        for i, el in enumerate(filhos):
            if i >= a_partir and el.tag == qn("w:p") and _texto_de(el).startswith(prefixo):
                return i
        raise ValueError(f"Parágrafo não encontrado: {prefixo!r}")

    # ---------------- cabeçalho institucional (tabela 0) e abertura
    celula = d.tables[0].rows[0].cells[2]
    pars_celula = [p for p in celula.paragraphs]
    for p in pars_celula:
        t = p.text.strip()
        if "REGIÃO DE POLÍCIA MILITAR" in t:
            _definir_texto(p, "{{ linha_regiao }}")
        elif "BATALHÃO DE POLÍCIA MILITAR" in t:
            _definir_texto(p, "{{ linha_unidade }}")

    for par in d.paragraphs[:20]:
        t = par.text.strip()
        if t.startswith("Quartel em"):
            _definir_texto(par, "Quartel em {{ cidade_sede }}, {{ data_documento_extenso }}.")
        elif t.startswith("Ao Senhor"):
            _definir_texto(par, "Ao Senhor {{ destinatario }}")
        elif t.startswith("Anexo:"):
            _definir_texto(par, "Anexo: {{ anexo_descricao }}")

    # ---------------- itens 1 e 2 (data e local do fato)
    i_data = indice_par("1. Data do Fato")
    _definir_texto(docx.text.paragraph.Paragraph(filhos[i_data + 1], d), "{{ data_fato_linha }}")
    i_local = indice_par("2. Local do Fato")
    _definir_texto(docx.text.paragraph.Paragraph(filhos[i_local + 1], d), "{{ local_fato_linha }}")

    # ---------------- descrição sucinta: vira laço de parágrafos
    i_desc = indice_par("2. Descrição sucinta")
    i_indiv = indice_par("3. Individualização")
    primeiro_texto = filhos[i_desc + 1]
    par_modelo = docx.text.paragraph.Paragraph(primeiro_texto, d)
    _definir_texto(par_modelo, "{{ par }}")
    vazio = _paragrafo_tag(par_modelo, "")
    corpo.insert(list(corpo).index(primeiro_texto), _paragrafo_tag(par_modelo, "{%p for par in descricao_paragrafos %}"))
    primeiro_texto.addnext(_paragrafo_tag(par_modelo, "{%p endfor %}"))
    primeiro_texto.addnext(vazio)
    # remove os demais parágrafos do exemplo entre a descrição e o item 3
    filhos = list(corpo)
    i_desc = indice_par("2. Descrição sucinta")
    i_fim = indice_par("{%p endfor")
    i_indiv = indice_par("3. Individualização")
    for el in filhos[i_fim + 1:i_indiv]:
        corpo.remove(el)

    # ---------------- individualização: um bloco em laço
    filhos = list(corpo)
    i_indiv = indice_par("3. Individualização")
    i_criterios = indice_par("4. Critérios")
    cabeca = filhos[i_indiv + 2] if _texto_de(filhos[i_indiv + 1]) == "" else filhos[i_indiv + 1]
    par_cabeca = docx.text.paragraph.Paragraph(cabeca, d)
    _definir_texto(
        par_cabeca,
        "3.{{ loop.index }}. Nº {{ m.numero }}, {{ m.posto_doc }} {{ m.nome_upper }}, {{ m.unidade }} / {{ m.funcao_upper }}.",
    )
    idx = list(corpo).index(cabeca)
    texto_el = list(corpo)[idx + 1]
    _definir_texto(docx.text.paragraph.Paragraph(texto_el, d), "{{ m.individualizacao }}")
    corpo.insert(idx, _paragrafo_tag(par_cabeca, "{%p for m in militares %}"))
    fim_vazio = _paragrafo_tag(par_cabeca, "")
    texto_el.addnext(_paragrafo_tag(par_cabeca, "{%p endfor %}"))
    texto_el.addnext(fim_vazio)
    filhos = list(corpo)
    i_fim = indice_par("{%p endfor", indice_par("3. Individualização"))
    i_criterios = indice_par("4. Critérios")
    for el in filhos[i_fim + 1:i_criterios]:
        corpo.remove(el)

    # ---------------- critérios: cabeçalho + tabela em laço
    filhos = list(corpo)
    i_criterios = indice_par("4. Critérios")
    i_41 = indice_par("4.1")
    cabeca41 = filhos[i_41]
    _definir_texto(
        docx.text.paragraph.Paragraph(cabeca41, d),
        "4.{{ loop.index }}. Nº {{ m.numero }}, {{ m.posto_doc }} {{ m.nome_upper }}, {{ m.unidade }} / {{ m.funcao_upper }}.",
    )
    # primeira tabela de requisitos depois do cabeçalho 4.1
    tabela1 = None
    for el in filhos[i_41:]:
        if el.tag == qn("w:tbl"):
            tabela1 = el
            break
    # marca as células Sim/Não (última célula das linhas 2-9)
    linhas = tabela1.findall(qn("w:tr"))
    for n, tr in enumerate(linhas[2:10], start=1):
        tc = tr.findall(qn("w:tc"))[-1]
        for t in tc.iter(qn("w:t")):
            t.text = ""
        primeiro_t = next(iter(tc.iter(qn("w:t"))), None)
        if primeiro_t is None:
            # célula sem run de texto: cria via parágrafo
            p = tc.find(qn("w:p"))
            r = p.makeelement(qn("w:r"), {})
            t = p.makeelement(qn("w:t"), {})
            t.text = f"{{{{ m.req{n} }}}}"
            r.append(t)
            p.append(r)
        else:
            primeiro_t.text = f"{{{{ m.req{n} }}}}"

    par_ref = docx.text.paragraph.Paragraph(cabeca41, d)
    corpo.insert(list(corpo).index(cabeca41), _paragrafo_tag(par_ref, "{%p for m in militares %}"))
    tabela1.addnext(_paragrafo_tag(par_ref, "{%p endfor %}"))
    tabela1.addnext(_paragrafo_tag(par_ref, ""))
    # remove os blocos 4.2-4.15 (cabeçalhos e tabelas) até a fundamentação
    filhos = list(corpo)
    i_fim = None
    for i, el in enumerate(filhos):
        if el.tag == qn("w:p") and _texto_de(el).startswith("{%p endfor") and i > indice_par("4. Critérios"):
            i_fim = i
    i_fund = indice_par("5. Fundamentação")
    for el in filhos[i_fim + 1:i_fund]:
        corpo.remove(el)

    # ---------------- parecer: tipo + lista de militares em laço
    for par in d.paragraphs:
        if par.text.strip().startswith("Sugiro a concessão"):
            _definir_texto(par, "Sugiro a concessão de {{ tipo_recompensa_upper }} para os militares:")
            break
    filhos = list(corpo)
    i_sugiro = indice_par("Sugiro a concessão")
    i_ass = None
    for i, el in enumerate(filhos):
        if el.tag == qn("w:p") and re.match(r"^[A-ZÀ-Ú ]+,\s", _texto_de(el)) and i > i_sugiro:
            i_ass = i  # linha de assinatura (NOME, POSTO.)
            break
    # a primeira linha de militar vira o corpo do laço; demais são removidas
    i_linha1 = None
    for i in range(i_sugiro + 1, i_ass):
        if _texto_de(filhos[i]).startswith(("Nº", "N°")):
            i_linha1 = i
            break
    linha1 = filhos[i_linha1]
    par_linha = docx.text.paragraph.Paragraph(linha1, d)
    _definir_texto(par_linha, "Nº {{ m.numero }}, {{ m.posto_doc }} {{ m.nome_upper }}, {{ m.unidade }} / {{ m.funcao_upper }}.")
    corpo.insert(list(corpo).index(linha1), _paragrafo_tag(par_linha, "{%p for m in militares %}"))
    linha1.addnext(_paragrafo_tag(par_linha, "{%p endfor %}"))
    filhos = list(corpo)
    i_fim = indice_par("{%p endfor", indice_par("Sugiro a concessão"))
    i_ass = None
    for i, el in enumerate(filhos):
        if el.tag == qn("w:p") and re.match(r"^[A-ZÀ-Ú ]+,\s", _texto_de(el)) and i > i_fim:
            i_ass = i
            break
    for el in filhos[i_fim + 1:i_ass]:
        if _texto_de(el).startswith(("Nº", "N°")):
            corpo.remove(el)

    # assinatura do interessado/proponente
    _definir_texto(docx.text.paragraph.Paragraph(filhos[i_ass], d), "{{ proponente_assinatura }}")

    # ---------------- anexos: remove os do caso e deixa laço de links
    filhos = list(corpo)
    i_rep = indice_par("REPORTAGENS")
    par_rep = docx.text.paragraph.Paragraph(filhos[i_rep], d)
    apos_rep = filhos[i_rep]
    apos_rep.addnext(_paragrafo_tag(par_rep, "{%p endfor %}"))
    apos_rep.addnext(_paragrafo_tag(par_rep, "{{ a }}"))
    apos_rep.addnext(_paragrafo_tag(par_rep, "{%p for a in anexos_linhas %}"))
    filhos = list(corpo)
    i_fim = indice_par("{%p endfor", indice_par("REPORTAGENS"))
    for el in filhos[i_fim + 1:]:
        if el.tag == qn("w:sectPr"):
            continue
        corpo.remove(el)

    d.save(str(destino))
    _remover_midia_orfa(destino)


def _remover_midia_orfa(caminho: Path) -> None:
    """Remove do pacote as imagens que não são mais referenciadas pelo corpo
    (as fotos dos anexos do caso de exemplo). O brasão, referenciado pelo
    documento, permanece."""
    origem = caminho.with_suffix(".tmp.docx")
    shutil.move(caminho, origem)
    with zipfile.ZipFile(origem) as z:
        doc_xml = z.read("word/document.xml").decode("utf-8")
        rels_xml = z.read("word/_rels/document.xml.rels").decode("utf-8")
        ids_usados = set(re.findall(r'r:(?:embed|id)="(rId\d+)"', doc_xml))
        remover_alvos = set()
        novas_rels = []
        for rel in re.findall(r"<Relationship [^>]*/>", rels_xml):
            rid = re.search(r'Id="(rId\d+)"', rel).group(1)
            alvo = re.search(r'Target="([^"]+)"', rel).group(1)
            if "image" in rel and rid not in ids_usados:
                remover_alvos.add("word/" + alvo.replace("../", ""))
                continue
            novas_rels.append(rel)
        rels_final = re.sub(r"<Relationship .*/>(?=</Relationships>)", "", rels_xml, flags=re.S)
        rels_final = re.sub(
            r"(<Relationships[^>]*>).*(</Relationships>)",
            lambda m: m.group(1) + "".join(novas_rels) + m.group(2),
            rels_xml, flags=re.S,
        )
        with zipfile.ZipFile(caminho, "w", zipfile.ZIP_DEFLATED) as saida:
            for item in z.infolist():
                if item.filename in remover_alvos:
                    continue
                dados = rels_final.encode("utf-8") if item.filename == "word/_rels/document.xml.rels" else z.read(item.filename)
                saida.writestr(item, dados)
    origem.unlink()


if __name__ == "__main__":
    construir(Path(sys.argv[1]))
    print(f"Template gerado em: {DESTINO}")
