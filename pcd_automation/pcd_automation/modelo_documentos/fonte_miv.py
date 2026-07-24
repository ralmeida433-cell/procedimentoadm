"""Padronização de fonte e layout conforme o Manual de Identidade Visual
(MIV) da PMMG (Res. nº 5.450/2025).

Fonte: a fonte oficial dos documentos institucionais passou a ser a
família **Rawline** (antes era Arial, pelo MEGDI/2019).
- Corpo do documento: Rawline, tamanho 11, cor preta.
- Cabeçalho de unidade (nome da instituição): Rawline Black, tamanho 11, preto.

Layout (regras de diagramação do MIV):
- Margens: 3 cm (superior e esquerda), 2 cm (direita e inferior).
- Entrelinhas: 1,5.
- Parágrafo moderno: sem recuo de início de parágrafo, com espaçamento de
  6 pt antes e depois.

Esta padronização é aplicada como pós-processamento no documento já
renderizado (após o merge do docxtpl), então NÃO interfere na fidelidade
de conteúdo nem no padrão de negrito herdado do modelo: só troca a família
tipográfica, garante cor preta, fixa margens/espaçamento e remove recuo de
início de parágrafo. NÃO mexe em alinhamento (centralizado/à direita/
justificado) nem em recuos NEGATIVOS (recuo suspenso usado em itens
enumerados como "I –", "II –") - esses fazem parte da estrutura do texto,
não do "recuo de parágrafo" que o MIV veda.

Importante: para o documento *exibir* de fato em Rawline, a fonte precisa
estar instalada na máquina que o abre. O .docx sempre passa a especificar
Rawline (que é o exigido pelo MIV); em máquinas sem a fonte, o Word faz a
substituição visual automaticamente.
"""
from __future__ import annotations

from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONTE_CORPO = "Rawline"
FONTE_CABECALHO = "Rawline Black"
TAMANHO = Pt(11)
PRETO = RGBColor(0x00, 0x00, 0x00)

# Marcadores para identificar a linha de cabeçalho da unidade (nome da
# instituição por extenso) - só esses trechos recebem Rawline Black.
MARCADORES_CABECALHO = ("POLÍCIA MILITAR DE MINAS GERAIS", "POLICIA MILITAR DE MINAS GERAIS")

MARGEM_SUPERIOR = Cm(3)
MARGEM_ESQUERDA = Cm(3)
MARGEM_DIREITA = Cm(2)
MARGEM_INFERIOR = Cm(2)
# Se o documento for assinado pelo Assinador Digital da Intranet PM, o MIV
# exige margem inferior de 6 cm - não aplicado por padrão (não há como o
# programa saber se o documento será assinado por esse meio específico).
MARGEM_INFERIOR_ASSINADOR_DIGITAL = Cm(6)

ENTRELINHAS = 1.5
ESPACO_ANTES = Pt(6)
ESPACO_DEPOIS = Pt(6)


def _definir_rfonts(rpr, nome: str) -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    for atributo in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(atributo), nome)


def _aplicar_run(run, nome: str) -> None:
    run.font.name = nome
    _definir_rfonts(run._element.get_or_add_rPr(), nome)
    run.font.color.rgb = PRETO
    if run.font.size is None:  # não encolhe título que já tenha tamanho próprio
        run.font.size = TAMANHO


def _e_cabecalho(paragrafo) -> bool:
    texto = (paragrafo.text or "").upper()
    return any(marcador in texto for marcador in MARCADORES_CABECALHO)


def _aplicar_paragrafos(paragrafos) -> None:
    for paragrafo in paragrafos:
        nome = FONTE_CABECALHO if _e_cabecalho(paragrafo) else FONTE_CORPO
        for run in paragrafo.runs:
            _aplicar_run(run, nome)


def _aplicar_tabelas(tabelas) -> None:
    for tabela in tabelas:
        for linha in tabela.rows:
            for celula in linha.cells:
                _aplicar_paragrafos(celula.paragraphs)
                _aplicar_tabelas(celula.tables)


def _aplicar_layout_paragrafo(paragrafo) -> None:
    pf = paragrafo.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = ENTRELINHAS
    pf.space_before = ESPACO_ANTES
    pf.space_after = ESPACO_DEPOIS
    # "Parágrafo moderno": sem recuo de início de parágrafo. Só remove
    # recuo POSITIVO (o recuo tradicional de começo de parágrafo) - recuos
    # negativos (recuo suspenso de itens enumerados "I –", "II –") e o
    # alinhamento do parágrafo (centralizado, à direita etc.) são
    # preservados, pois fazem parte da estrutura do texto do modelo.
    recuo = pf.first_line_indent
    if recuo is not None and recuo.pt > 0:
        pf.first_line_indent = None


def _aplicar_layout_paragrafos(paragrafos) -> None:
    for paragrafo in paragrafos:
        _aplicar_layout_paragrafo(paragrafo)


def _aplicar_layout_tabelas(tabelas) -> None:
    for tabela in tabelas:
        for linha in tabela.rows:
            for celula in linha.cells:
                _aplicar_layout_paragrafos(celula.paragraphs)
                _aplicar_layout_tabelas(celula.tables)


def aplicar_layout_miv(doc) -> None:
    """Aplica margens, entrelinhas e parágrafo moderno do MIV.

    Margens: 3 cm (superior/esquerda), 2 cm (direita/inferior). Entrelinhas
    1,5. Espaçamento de 6 pt antes/depois de cada parágrafo do corpo
    (cabeçalho e rodapé de seção não são alterados - são elementos curtos
    de identidade visual, não "corpo do documento").
    """
    for secao in doc.sections:
        secao.top_margin = MARGEM_SUPERIOR
        secao.left_margin = MARGEM_ESQUERDA
        secao.right_margin = MARGEM_DIREITA
        secao.bottom_margin = MARGEM_INFERIOR

    _aplicar_layout_paragrafos(doc.paragraphs)
    _aplicar_layout_tabelas(doc.tables)


def aplicar_fonte_miv(doc) -> None:
    """Aplica a fonte do MIV (Rawline) em todo o documento python-docx."""
    # Estilo Normal como base: cobre os runs que não têm fonte explícita
    # (que, sem isso, cairiam no padrão do Word - Calibri/Times - e não em Rawline).
    try:
        normal = doc.styles["Normal"]
        normal.font.name = FONTE_CORPO
        normal.font.size = TAMANHO
        normal.font.color.rgb = PRETO
        _definir_rfonts(normal.element.get_or_add_rPr(), FONTE_CORPO)
    except Exception:
        pass  # se o modelo não tiver estilo Normal, os runs ainda recebem a fonte abaixo

    _aplicar_paragrafos(doc.paragraphs)
    _aplicar_tabelas(doc.tables)
    for secao in doc.sections:
        for parte in (secao.header, secao.footer, secao.first_page_header, secao.first_page_footer):
            _aplicar_paragrafos(parte.paragraphs)
            _aplicar_tabelas(parte.tables)


def aplicar_padrao_miv(doc) -> None:
    """Aplica o padrão completo do MIV: fonte Rawline + layout (margens,
    entrelinhas, parágrafo moderno). Função única a ser chamada após o
    preenchimento do documento."""
    aplicar_fonte_miv(doc)
    aplicar_layout_miv(doc)
