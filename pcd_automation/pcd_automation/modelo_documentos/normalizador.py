"""Motor de normalização dos modelos .docx originais.

Os modelos usam preenchimento manual (ex.: "XX", "FULANO DE TAL") em que o
autor do modelo **negritou especificamente os trechos a preencher**,
deixando o texto fixo (boilerplate) sem negrito - essa é a convenção visual
do próprio modelo, não uma invenção nossa (confirmado inspecionando os
`.runs` de vários documentos: o rótulo e o texto de conexão estão em runs
não-negritados, e cada branco está isolado no seu próprio run negritado).

Por isso a normalização troca o texto **apenas dos runs que já eram os
blancos** por marcadores Jinja2 (`{{ campo }}`), sem tocar nos demais runs.
Isso preserva 100% da formatação original do documento - o resultado final,
depois do merge com `docxtpl` (ver `preenchedor.py`), fica visualmente
idêntico ao modelo: rótulos e texto fixo do jeito que já estavam, e só os
dados preenchidos aparecem em negrito, exatamente como no modelo em branco.

Substituir o parágrafo inteiro (abordagem anterior) forçava colapsar todos
os runs em um só, herdando a formatação de apenas um deles para o
parágrafo inteiro - o que ora deixava frases inteiras em negrito indevido,
ora apagava o negrito que deveria destacar o dado preenchido. Este módulo
substitui aquela abordagem.
"""
from __future__ import annotations

from collections import defaultdict, deque
from pathlib import Path

import docx

# (índice_coluna_da_célula, índice_parágrafo_na_célula, texto_original, texto_novo)
SubstituicaoTabela = tuple[int, int, str, str]
# (texto_original_do_run, texto_novo_do_run_com_marcador_jinja)
SubstituicaoRun = tuple[str, str]
# (texto_completo_original_do_parágrafo - usado só para localizar o parágrafo certo,
#  lista de substituições de run dentro dele)
SubstituicaoParagrafo = tuple[str, list[SubstituicaoRun]]


def _substituir_texto_completo(paragrafo, novo_texto: str) -> None:
    """Substitui todo o texto de um parágrafo, mantendo a formatação do
    primeiro run - usado só para células de tabela/timbre, que têm um único
    estilo uniforme (não há negrito seletivo a preservar ali)."""
    if paragrafo.runs:
        paragrafo.runs[0].text = novo_texto
        for run in paragrafo.runs[1:]:
            run.text = ""
    else:
        paragrafo.add_run(novo_texto)


def _substituir_runs(paragrafo, substituicoes: list[SubstituicaoRun]) -> set[str]:
    """Troca o texto de runs específicos, preservando a formatação de cada
    um. Quando o mesmo texto de run se repete no parágrafo (ex.: "XX"
    aparecendo mais de uma vez com significados diferentes), cada
    ocorrência consome a próxima substituição da fila, na ordem em que
    aparecem no parágrafo - por isso a ORDEM das entradas repetidas em
    `substituicoes` deve seguir a ordem em que os runs aparecem no modelo.

    Retorna o conjunto de textos de run que não foram encontrados (para o
    chamador validar que nada ficou faltando).
    """
    fila: dict[str, deque[str]] = defaultdict(deque)
    for original, novo in substituicoes:
        fila[original].append(novo)

    for run in paragrafo.runs:
        pendentes = fila.get(run.text)
        if pendentes:
            run.text = pendentes.popleft()

    return {original for original, restantes in fila.items() if restantes}


def gerar_normalizado(
    origem: Path,
    destino: Path,
    substituicoes_paragrafo: list[SubstituicaoParagrafo],
    substituicoes_tabela: list[SubstituicaoTabela] | None = None,
) -> Path:
    if not origem.exists():
        raise FileNotFoundError(f"Modelo original não encontrado: {origem}")

    documento = docx.Document(str(origem))

    fila_paragrafos: dict[str, deque[list[SubstituicaoRun]]] = defaultdict(deque)
    for texto_ancora, subs_runs in substituicoes_paragrafo:
        fila_paragrafos[texto_ancora].append(subs_runs)

    for paragrafo in documento.paragraphs:
        fila = fila_paragrafos.get(paragrafo.text)
        if not fila:
            continue
        subs_runs = fila.popleft()
        faltantes = _substituir_runs(paragrafo, subs_runs)
        if faltantes:
            raise ValueError(
                f"Run(s) esperado(s) não encontrado(s) no parágrafo {paragrafo.text!r} "
                f"de {origem.name} (o arquivo pode ter mudado): {faltantes}"
            )

    paragrafos_restantes = [texto for texto, fila in fila_paragrafos.items() if fila]
    if paragrafos_restantes:
        raise ValueError(
            f"Parágrafo(s) esperado(s) não encontrado(s) (ou com menos ocorrências que o "
            f"esperado) em {origem.name} (o arquivo pode ter mudado): {paragrafos_restantes}"
        )

    for coluna_idx, paragrafo_idx, original, novo in substituicoes_tabela or []:
        celula = documento.tables[0].rows[0].cells[coluna_idx]
        paragrafo = celula.paragraphs[paragrafo_idx]
        if paragrafo.text != original:
            raise ValueError(
                f"Parágrafo da tabela [{coluna_idx}][{paragrafo_idx}] em {origem.name} não confere: "
                f"esperado {original!r}, encontrado {paragrafo.text!r}"
            )
        _substituir_texto_completo(paragrafo, novo)

    destino.parent.mkdir(parents=True, exist_ok=True)
    documento.save(str(destino))
    return destino
