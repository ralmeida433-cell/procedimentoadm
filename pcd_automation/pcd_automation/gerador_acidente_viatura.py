"""Geração das peças da Sindicância de Acidente com Viatura.

Base normativa (MAPPA, Capítulo IX - arts. 321 a 324):

- art. 321: o acidente com viatura é apurado por SAD, que verifica a eventual
  transgressão disciplinar, o dano causado à viatura e a responsabilidade pelo
  ressarcimento ao erário ou ao particular;
- art. 322, I a IV: define quando, ALÉM da SAD, é obrigatório instaurar IPM
  (regra geral: sempre que houver vítima). Ver `alertas_ipm`;
- art. 324: aplicam-se à sindicância de acidente com viatura "as normas e os
  modelos contidos no capítulo específico sobre SAD" - por isso as peças aqui
  usam os modelos oficiais da SAD como base visual (cabeçalho, brasão,
  estilos), em vez de um layout inventado;
- art. 324, §2º (c/c art. 72 da ICCPM/BM 01/14): quando a sindicância apura
  EXCLUSIVAMENTE a responsabilidade pelos danos, sem caráter disciplinar (ex.:
  culpa exclusiva de terceiro civil), ela "poderá ser elaborada integralmente
  sem notificação do militar sindicado para defesa prévia", indo do termo de
  abertura direto à colheita de provas. Ver `PECAS` e `pecas_aplicaveis`.

Prazo: 30 dias corridos, prorrogáveis por 10 (art. 273/274, aplicados por
força do art. 324).

Nota sobre nomenclatura: o pedido original chamava este procedimento de
"PAAV". O termo não existe no MAPPA - a busca no texto integral do manual
retorna zero ocorrências. O instrumento correto para acidente com viatura é a
SAD (art. 321), e é essa a nomenclatura usada nas peças, porque documento com
sigla inexistente é vício formal.

Três das sete peças (Portaria, Termo de Encerramento e Solução) não têm modelo
oficial na pasta de modelos da PMMG deste projeto; são montadas aqui a partir
do mesmo cabeçalho institucional das demais, seguindo a redação oficial.
"""
from __future__ import annotations

import copy
from dataclasses import dataclass, field
from datetime import date, timedelta
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .modelo_documentos.cabecalho_layout import padronizar_brasao_esquerda, padronizar_carimbo_direita
from .modelo_documentos.fonte_miv import aplicar_padrao_miv
from .modelo_documentos.formatacao import MESES_PT, data_por_extenso
from .redacao import formatar_hora_br

# Modelo oficial da SAD usado só como BASE VISUAL: dele vêm o cabeçalho com o
# brasão, os estilos e as margens. O corpo é sempre limpo e reescrito.
CAMINHO_BASE = (
    Path(__file__).resolve().parents[2]
    / "Modelo Processo-Procedimento" / "SAD" / "02 Termo de Abertura.docx"
)

PRAZO_DIAS = 30            # art. 273, aplicado por força do art. 324
PRAZO_PRORROGACAO = 10     # art. 274

# (id, nome do arquivo, título da peça). A ordem é a de juntada nos autos.
PECAS: list[tuple[str, str, str]] = [
    ("portaria", "01_portaria_instauracao.docx", "Portaria de Instauração"),
    ("abertura", "02_termo_abertura.docx", "Termo de Abertura dos Trabalhos"),
    ("notificacao", "03_notificacao_sindicado.docx", "Notificação de Instauração e Citação"),
    ("depoimento", "04_termo_depoimento.docx", "Termo de Depoimento / Inquirição"),
    ("relatorio", "05_relatorio_final.docx", "Relatório Final do Encarregado"),
    ("encerramento", "06_termo_encerramento.docx", "Termo de Encerramento dos Trabalhos"),
    ("solucao", "07_solucao.docx", "Solução da Sindicância"),
]


@dataclass
class ResultadoAcidenteViatura:
    ok: bool = True
    documentos: list[str] = field(default_factory=list)
    pendentes: list[str] = field(default_factory=list)
    alertas: list[str] = field(default_factory=list)
    erros: list[str] = field(default_factory=list)


# --------------------------------------------------------------- regras MAPPA

def pecas_aplicaveis(dados: dict) -> list[tuple[str, str, str]]:
    """Peças que se aplicam ao caso.

    Sem caráter disciplinar (art. 324, §2º), a notificação do sindicado para
    defesa prévia é dispensada - a sindicância vai do termo de abertura direto
    à colheita de provas.
    """
    if _sim(dados.get("sem_carater_disciplinar")):
        return [p for p in PECAS if p[0] != "notificacao"]
    return list(PECAS)


def alertas_ipm(dados: dict) -> list[str]:
    """Avisa quando o art. 322 exige instaurar IPM ALÉM da sindicância.

    Não instaurar o IPM devido é omissão grave: a apuração penal do fato ficaria
    sem instrumento. O sistema não instaura nada sozinho - só avisa.
    """
    avisos: list[str] = []
    vitima = (dados.get("houve_vitima") or "").strip().lower()
    if vitima in ("", "nao", "não", "n", "sem vitima", "sem vítima"):
        return avisos

    if "civil" in vitima:
        avisos.append(
            "Acidente com vítima civil: além desta sindicância, instaure IPM para apurar a infração "
            "penal (art. 322, I, do MAPPA). O IPM segue para a JME sem análise de transgressões "
            "residuais - é esta SAD que cuida do aspecto disciplinar e do ressarcimento."
        )
    elif "condutor" in vitima:
        avisos.append(
            "Vítima militar que conduzia a viatura: NÃO há crime militar de autolesão, então basta a "
            "sindicância para apurar a responsabilidade civil e disciplinar (art. 322, III, do MAPPA). "
            "Se houver indício de crime comum praticado por civil causador, encaminhe o fato via "
            "Boletim de Ocorrência à autoridade de polícia judiciária e registre isso nos autos. "
            "Havendo suspeita de sabotagem, ou lesão decorrente de ordem de superior, instaure IPM."
        )
    elif "militar" in vitima:
        avisos.append(
            "Acidente com vítima militar (passageiro ou pedestre): além desta sindicância, instaure "
            "IPM para apurar a infração penal (art. 322, II, do MAPPA)."
        )

    if avisos:
        avisos.append(
            "Havendo IPM e SAD sobre o mesmo acidente, o encargo de sindicante e o de encarregado do "
            "IPM devem recair, preferencialmente, sobre o mesmo militar (art. 322, parágrafo único)."
        )
    return avisos


def prazo_conclusao(data_recibo: date | None, prorrogado: bool = False) -> date | None:
    """Termo final do prazo, contado do recibo da portaria pelo sindicante
    (art. 273/274 c/c art. 324)."""
    if not data_recibo:
        return None
    dias = PRAZO_DIAS + (PRAZO_PRORROGACAO if prorrogado else 0)
    return data_recibo + timedelta(days=dias)


# --------------------------------------------------------------- utilitários

def _sim(valor) -> bool:
    if isinstance(valor, bool):
        return valor
    return str(valor or "").strip().lower() in ("s", "sim", "true", "1")


def _texto(dados: dict, chave: str, rotulo: str, pendentes: list[str]) -> str:
    valor = str(dados.get(chave) or "").strip()
    if valor:
        return valor
    if rotulo not in pendentes:
        pendentes.append(rotulo)
    return f"[PREENCHER: {rotulo}]"


def _data(dados: dict, chave: str) -> date | None:
    valor = dados.get(chave)
    if isinstance(valor, date):
        return valor
    if isinstance(valor, str) and valor.strip():
        try:
            return date.fromisoformat(valor.strip())
        except ValueError:
            return None
    return None


def _data_barra(d: date | None, rotulo: str, pendentes: list[str]) -> str:
    if d is None:
        if rotulo not in pendentes:
            pendentes.append(rotulo)
        return f"[PREENCHER: {rotulo}]"
    return f"{d.day:02d}/{d.month:02d}/{d.year}"


def _qualificar(dados: dict, prefixo: str, rotulo: str, pendentes: list[str]) -> str:
    """Qualificação no corpo do texto, no formato dos modelos oficiais da PMMG:
    "nº 123.456-7, 1º Sargento PM FULANO DE TAL" - sem vírgula entre o posto e
    o nome, como aparece nos modelos da SAD."""
    numero = str(dados.get(f"numero_{prefixo}") or "").strip()
    posto = str(dados.get(f"posto_{prefixo}") or "").strip()
    nome = str(dados.get(f"nome_{prefixo}") or "").strip()
    if not nome:
        if rotulo not in pendentes:
            pendentes.append(rotulo)
        return f"[PREENCHER: {rotulo}]"
    identificacao = f"{posto} PM {nome.upper()}" if posto else nome.upper()
    return f"nº {numero}, {identificacao}" if numero else identificacao


def _para_assinatura(dados: dict, prefixo: str) -> str:
    """Bloco de assinatura: "FULANO DE TAL, CAPITÃO PM" - ordem invertida em
    relação ao corpo do texto, seguindo os modelos oficiais."""
    posto = str(dados.get(f"posto_{prefixo}") or "").strip()
    nome = str(dados.get(f"nome_{prefixo}") or "").strip()
    if not nome:
        return "[PREENCHER: nome para assinatura]"
    return f"{nome.upper()}, {posto.upper()} PM" if posto else nome.upper()


# --------------------------------------------------------------- montagem docx

def _abrir_base():
    """Abre o modelo oficial da SAD e limpa o corpo, preservando a tabela do
    cabeçalho (que contém o brasão) e todo o restante da formatação."""
    if not CAMINHO_BASE.exists():
        raise FileNotFoundError(
            f"Modelo oficial da SAD não encontrado em {CAMINHO_BASE}. Ele é a base visual "
            "(cabeçalho e brasão) das peças da sindicância de acidente com viatura."
        )
    doc = docx.Document(str(CAMINHO_BASE))
    corpo = doc.element.body
    # Remove tudo depois da PRIMEIRA tabela (o cabeçalho institucional).
    tabela_vista = False
    for filho in list(corpo):
        marcador = filho.tag.split("}")[-1]
        if marcador == "tbl" and not tabela_vista:
            tabela_vista = True
            continue
        if marcador == "sectPr":
            continue
        if tabela_vista:
            corpo.remove(filho)
    return doc


def _p(doc, texto: str = "", *, negrito=False, centro=False, tamanho=11,
       maiusculas=False, recuo=True, espaco_antes=0):
    par = doc.add_paragraph()
    if centro:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif texto:
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if recuo:
            par.paragraph_format.first_line_indent = Pt(35)
    par.paragraph_format.space_before = Pt(espaco_antes)
    par.paragraph_format.space_after = Pt(6)
    run = par.add_run(texto.upper() if maiusculas else texto)
    run.bold = negrito
    run.font.size = Pt(tamanho)
    return par


def _assinatura(doc, nome_cargo: str, cargo: str):
    _p(doc, "", recuo=False)
    _p(doc, nome_cargo.upper(), centro=True, negrito=True)
    _p(doc, cargo.upper(), centro=True)


def _fecho_local_data(doc, cidade: str, quando: date | None):
    quando = quando or date.today()
    _p(doc, f"Quartel em {cidade}, {data_por_extenso(quando)}.", centro=True, recuo=False, espaco_antes=12)


def _salvar(doc, caminho: Path) -> Path:
    caminho.parent.mkdir(parents=True, exist_ok=True)
    aplicar_padrao_miv(doc)
    padronizar_brasao_esquerda(doc)
    doc.save(str(caminho))
    padronizar_carimbo_direita(caminho)
    return caminho


# --------------------------------------------------------------- as sete peças

def _ctx(dados: dict, pendentes: list[str]) -> dict:
    cidade = _texto(dados, "cidade_sede", "cidade do quartel", pendentes)
    return {
        "cidade": cidade,
        "unidade": _texto(dados, "unidade", "unidade/UDI", pendentes),
        "portaria": _texto(dados, "numero_portaria", "número da portaria", pendentes),
        "data_portaria": _data(dados, "data_portaria"),
        "sindicante": _qualificar(dados, "sindicante", "dados do sindicante (encarregado)", pendentes),
        "sindicante_assina": _para_assinatura(dados, "sindicante"),
        "delegante": _qualificar(dados, "delegante", "dados da autoridade delegante", pendentes),
        "delegante_assina": _para_assinatura(dados, "delegante"),
        "cargo_delegante": _texto(dados, "cargo_delegante", "cargo da autoridade delegante", pendentes),
        "envolvido": _qualificar(dados, "envolvido", "dados do envolvido/responsável", pendentes),
        "viatura": _texto(dados, "descricao_viatura", "descrição da viatura/material", pendentes),
        "prejuizo": str(dados.get("valor_prejuizo") or "").strip(),
        "data_fato": _data(dados, "data_fato"),
        "hora_fato": formatar_hora_br(dados.get("hora_fato")) if dados.get("hora_fato") else "",
        "local_fato": _texto(dados, "local_fato", "local do fato", pendentes),
        "historico": _texto(dados, "historico_fato", "histórico sucinto do fato", pendentes),
        "sem_disciplinar": _sim(dados.get("sem_carater_disciplinar")),
    }


def _peca_portaria(doc, c, dados, pendentes):
    _p(doc, f"PORTARIA Nº {c['portaria']}", centro=True, negrito=True, recuo=False)
    _p(doc, "INSTAURA SINDICÂNCIA ADMINISTRATIVA DISCIPLINAR PARA APURAÇÃO DE ACIDENTE COM VIATURA",
       centro=True, negrito=True, recuo=False)
    _p(doc, "")
    _p(doc, f"O {c['cargo_delegante']}, no uso das atribuições que lhe confere a legislação em vigor, "
            f"com fundamento no art. 321 e seguintes do MAPPA (Resolução Conjunta nº 4.220/2012) e na "
            f"Lei Estadual nº 14.310/2002 (CEDM), RESOLVE:")
    _p(doc, f"Art. 1º Instaurar Sindicância Administrativa Disciplinar para apurar o acidente ocorrido "
            f"em {_data_barra(c['data_fato'], 'data do fato', pendentes)}"
            f"{', por volta das ' + c['hora_fato'] if c['hora_fato'] else ''}, "
            f"em {c['local_fato']}, envolvendo {c['viatura']}, bem como o dano causado, a eventual "
            f"transgressão disciplinar e a responsabilidade pelo ressarcimento ao erário ou ao "
            f"particular, nos termos do art. 321 do MAPPA.")
    _p(doc, f"Art. 2º Designar o {c['sindicante']} para funcionar como sindicante.")
    _p(doc, f"Art. 3º Fixar o prazo de {PRAZO_DIAS} (trinta) dias corridos para a conclusão dos "
            f"trabalhos, contados do recibo desta portaria pelo sindicante, prorrogável por até "
            f"{PRAZO_PRORROGACAO} (dez) dias, mediante justificativa fundamentada nos autos "
            f"(arts. 273 e 274 do MAPPA).")
    _p(doc, "Art. 4º Esta portaria entra em vigor na data de sua publicação.")
    _fecho_local_data(doc, c["cidade"], c["data_portaria"])
    _assinatura(doc, c["delegante_assina"], c["cargo_delegante"])


def _peca_abertura(doc, c, dados, pendentes):
    _p(doc, "TERMO DE ABERTURA DOS TRABALHOS", centro=True, negrito=True, recuo=False)
    _p(doc, "")
    data_abertura = _data(dados, "data_abertura") or c["data_portaria"]
    quando = data_abertura or date.today()
    _p(doc, f"Aos {quando.day:02d} dias do mês de {MESES_PT[quando.month]} do ano de {quando.year}, "
            f"nesta cidade de {c['cidade']}, no cartório do(a) {c['unidade']}, em cumprimento à "
            f"Portaria nº {c['portaria']}, dei início a esta sindicância, destinada a apurar o "
            f"acidente com a viatura descrita nos autos, procedendo aos levantamentos e às entrevistas "
            f"preliminares e fazendo juntar aos autos os documentos a seguir relacionados: "
            f"{_texto(dados, 'documentos_juntados', 'documentos juntados na abertura', pendentes)}.")
    if c["sem_disciplinar"]:
        _p(doc, "Registro que a presente sindicância se destina exclusivamente a apurar a "
                "responsabilidade pelos danos causados, sem caráter disciplinar, razão pela qual foi "
                "dispensada a notificação de militar sindicado para apresentação de defesa prévia, na "
                "forma do art. 324, §2º, do MAPPA, c/c art. 72 da ICCPM/BM nº 01/2014.")
    else:
        _p(doc, "Foi procedida, ainda, a notificação do sindicado para que tomasse conhecimento da "
                "instauração do presente procedimento e apresentasse sua Defesa Prévia, conforme "
                "adiante se vê.")
    _p(doc, "Do que, para constar, lavro e assino o presente termo.")
    _fecho_local_data(doc, c["cidade"], data_abertura)
    _assinatura(doc, c["sindicante_assina"], "Sindicante")


def _peca_notificacao(doc, c, dados, pendentes):
    _p(doc, f"PORTARIA Nº {c['portaria']}", centro=True, negrito=True, recuo=False)
    _p(doc, "TERMO DE NOTIFICAÇÃO DE MILITAR SINDICADO", centro=True, negrito=True, recuo=False)
    _p(doc, "OBJETO DA ACUSAÇÃO E APRESENTAÇÃO DE DEFESA PRÉVIA", centro=True, negrito=True, recuo=False)
    _p(doc, "")
    _p(doc, f"Ao {c['envolvido']}", recuo=False)
    _p(doc, f"Anexos: autos da Portaria nº {c['portaria']}, contendo "
            f"{_texto(dados, 'numero_folhas_autos', 'nº de folhas dos autos', pendentes)} folhas.", recuo=False)
    _p(doc, "")
    _p(doc, f"Notifico-lhe que, em razão da instauração da Portaria nº {c['portaria']}, apura-se nesta "
            f"sindicância o acidente ocorrido em {_data_barra(c['data_fato'], 'data do fato', pendentes)}"
            f"{', por volta das ' + c['hora_fato'] if c['hora_fato'] else ''}, em {c['local_fato']}, "
            f"envolvendo {c['viatura']}, apurando-se o dano causado, a eventual transgressão "
            f"disciplinar e a responsabilidade pelo ressarcimento (art. 321 do MAPPA). Síntese do fato: "
            f"{c['historico']}")
    _p(doc, "Em razão das diligências que serão realizadas no processo, faculto-lhe acompanhar "
            "pessoalmente ou por defensor devidamente constituído (militar estadual de maior "
            "precedência hierárquica ou advogado) todos os atos a serem praticados.")
    _p(doc, "O rol de testemunhas de defesa, a inclusão de documentos e a produção de provas de "
            "interesse da defesa, caso não sejam apresentados na defesa prévia, poderão ser "
            "apresentados durante a instrução, até a abertura de vista para a defesa final.")
    _p(doc, "Fica ciente, ainda, de que ao final da instrução, caso reste alguma acusação contra a sua "
            "pessoa, ser-lhe-á dada nova vista dos autos (TAV) para a apresentação das Razões Escritas "
            "de Defesa finais (RED).")
    _p(doc, "A apresentação da Defesa Prévia é facultativa e deve ocorrer no prazo de 2 (dois) dias "
            "úteis, podendo o sindicado apresentá-la por ocasião de seu interrogatório (art. 291 do "
            "MAPPA). Esse prazo não é computado no prazo regulamentar de conclusão.")
    _p(doc, "")
    _p(doc, "RECEBI a presente NOTIFICAÇÃO e a documentação citada no anexo, e estou ciente sobre a "
            "faculdade de apresentar a defesa prévia, o rol de testemunhas e as provas que julgar "
            "necessárias.", recuo=False)
    _fecho_local_data(doc, c["cidade"], _data(dados, "data_notificacao"))
    _p(doc, "")
    _p(doc, "_________________________________________", centro=True, recuo=False)
    _p(doc, "SINDICADO", centro=True, recuo=False)
    _p(doc, "")
    _p(doc, "_________________________________________", centro=True, recuo=False)
    _p(doc, "SINDICANTE", centro=True, recuo=False)


def _peca_depoimento(doc, c, dados, pendentes):
    _p(doc, "TERMO DE DEPOIMENTO", centro=True, negrito=True, recuo=False)
    _p(doc, "")
    data_dep = _data(dados, "data_depoimento") or date.today()
    hora_dep = formatar_hora_br(dados.get("hora_depoimento")) if dados.get("hora_depoimento") else ""
    depoente = _qualificar(dados, "depoente", "dados do depoente", pendentes)
    _p(doc, f"Aos {data_dep.day:02d} dias do mês de {MESES_PT[data_dep.month]} do ano de "
            f"{data_dep.year}{', às ' + hora_dep if hora_dep else ''}, nesta cidade de {c['cidade']}, "
            f"no cartório do(a) {c['unidade']}, perante mim, {c['sindicante']}, sindicante designado "
            f"pela Portaria nº {c['portaria']}, presente o(a) depoente {depoente}, a quem foi "
            f"perguntado sobre os fatos apurados nesta sindicância, respondeu:")
    _p(doc, "")
    _p(doc, "QUANTO ÀS CIRCUNSTÂNCIAS DO FATO:", negrito=True, recuo=False)
    _p(doc, _texto(dados, "teor_circunstancias",
                   "teor do depoimento quanto às circunstâncias do fato", pendentes))
    _p(doc, "QUANTO À CAUTELA E AO ESTADO DE CONSERVAÇÃO DA VIATURA:", negrito=True, recuo=False)
    _p(doc, _texto(dados, "teor_cautela",
                   "teor do depoimento quanto à cautela e conservação da viatura", pendentes))
    _p(doc, "QUANTO AO NEXO DE CAUSALIDADE:", negrito=True, recuo=False)
    _p(doc, _texto(dados, "teor_nexo", "teor do depoimento quanto ao nexo de causalidade", pendentes))
    _p(doc, "")
    _p(doc, "Nada mais havendo, foi encerrado o presente termo, que, lido e achado conforme, vai "
            "assinado pelo depoente e por mim, sindicante.")
    _fecho_local_data(doc, c["cidade"], data_dep)
    _p(doc, "")
    _p(doc, "_________________________________________", centro=True, recuo=False)
    _p(doc, "DEPOENTE", centro=True, recuo=False)
    _p(doc, "")
    _p(doc, "_________________________________________", centro=True, recuo=False)
    _p(doc, "SINDICANTE", centro=True, recuo=False)


def _peca_relatorio(doc, c, dados, pendentes):
    _p(doc, "RELATÓRIO FINAL", centro=True, negrito=True, recuo=False)
    _p(doc, "")
    _p(doc, "1. PRELIMINARES E HISTÓRICO DOS FATOS", negrito=True, recuo=False)
    _p(doc, f"a. Portaria nº {c['portaria']}, de "
            f"{_data_barra(c['data_portaria'], 'data da portaria', pendentes)}.")
    _p(doc, f"b. Sindicante: {c['sindicante']}.")
    _p(doc, f"c. Envolvido/responsável pelo material: {c['envolvido']}.")
    _p(doc, f"d. Viatura/material: {c['viatura']}.")
    _p(doc, f"e. Local e data do fato: {c['local_fato']}, em "
            f"{_data_barra(c['data_fato'], 'data do fato', pendentes)}"
            f"{', às ' + c['hora_fato'] if c['hora_fato'] else ''}.")
    if c["prejuizo"]:
        _p(doc, f"f. Valor do prejuízo/avaria apurado: {c['prejuizo']}.")
    _p(doc, f"g. Histórico: {c['historico']}")
    _p(doc, "")
    _p(doc, "2. DILIGÊNCIAS REALIZADAS", negrito=True, recuo=False)
    _p(doc, _texto(dados, "diligencias", "diligências realizadas (seção 2 do relatório)", pendentes))
    _p(doc, "")
    _p(doc, "3. ANÁLISE DO MÉRITO", negrito=True, recuo=False)
    _p(doc, _texto(dados, "analise_merito",
                   "análise do mérito - culpa, dolo, caso fortuito, força maior ou desgaste natural "
                   "(seção 3 do relatório)", pendentes))
    _p(doc, "")
    _p(doc, "4. CONCLUSÃO E PARECER", negrito=True, recuo=False)
    _p(doc, _texto(dados, "conclusao_parecer",
                   "conclusão e parecer sobre a responsabilidade e o ressarcimento "
                   "(seção 4 do relatório)", pendentes))
    _fecho_local_data(doc, c["cidade"], _data(dados, "data_relatorio"))
    _assinatura(doc, c["sindicante_assina"], "Sindicante")


def _peca_encerramento(doc, c, dados, pendentes):
    _p(doc, "TERMO DE ENCERRAMENTO DOS TRABALHOS", centro=True, negrito=True, recuo=False)
    _p(doc, "")
    quando = _data(dados, "data_encerramento") or _data(dados, "data_relatorio") or date.today()
    _p(doc, f"Aos {quando.day:02d} dias do mês de {MESES_PT[quando.month]} do ano de {quando.year}, "
            f"nesta cidade de {c['cidade']}, no cartório do(a) {c['unidade']}, dou por encerrados os "
            f"trabalhos da sindicância instaurada pela Portaria nº {c['portaria']}, destinada a apurar "
            f"o acidente com viatura de que tratam os autos, os quais seguem com "
            f"{_texto(dados, 'numero_folhas_autos_final', 'nº total de folhas dos autos', pendentes)} "
            f"folhas, devidamente numeradas e rubricadas, para remessa à autoridade delegante.")
    _p(doc, "Do que, para constar, lavro e assino o presente termo.")
    _fecho_local_data(doc, c["cidade"], quando)
    _assinatura(doc, c["sindicante_assina"], "Sindicante")


def _peca_solucao(doc, c, dados, pendentes):
    _p(doc, "SOLUÇÃO", centro=True, negrito=True, recuo=False)
    _p(doc, f"Sindicância instaurada pela Portaria nº {c['portaria']}", centro=True, recuo=False)
    _p(doc, "")
    _p(doc, f"Trata-se de Sindicância Administrativa Disciplinar instaurada para apurar o acidente com "
            f"viatura ocorrido em {_data_barra(c['data_fato'], 'data do fato', pendentes)}, em "
            f"{c['local_fato']}, envolvendo {c['viatura']}, na forma do art. 321 do MAPPA.")
    _p(doc, "Examinados os autos e o relatório do sindicante, DECIDO:")
    _p(doc, _texto(dados, "texto_solucao",
                   "teor da decisão da autoridade delegante (homologar, reformar ou determinar "
                   "novas diligências, com fundamentação)", pendentes))
    _p(doc, f"Encaminhamentos: {_texto(dados, 'encaminhamentos', 'encaminhamentos administrativos e financeiros', pendentes)}")
    _p(doc, "Publique-se e cumpra-se.")
    _fecho_local_data(doc, c["cidade"], _data(dados, "data_solucao"))
    _assinatura(doc, c["delegante_assina"], c["cargo_delegante"])


_MONTADORES = {
    "portaria": _peca_portaria,
    "abertura": _peca_abertura,
    "notificacao": _peca_notificacao,
    "depoimento": _peca_depoimento,
    "relatorio": _peca_relatorio,
    "encerramento": _peca_encerramento,
    "solucao": _peca_solucao,
}


def gerar_documentos(dados: dict, diretorio_saida: Path) -> ResultadoAcidenteViatura:
    """Gera as peças aplicáveis ao caso em `diretorio_saida`."""
    resultado = ResultadoAcidenteViatura()
    pendentes: list[str] = []
    try:
        contexto = _ctx(dados, pendentes)
    except Exception as exc:  # noqa: BLE001 - erro vira mensagem para o usuário
        resultado.ok = False
        resultado.erros.append(f"Falha ao preparar os dados: {exc}")
        return resultado

    for peca_id, nome_arquivo, _titulo in pecas_aplicaveis(dados):
        try:
            doc = _abrir_base()
            _MONTADORES[peca_id](doc, contexto, dados, pendentes)
            _salvar(doc, Path(diretorio_saida) / nome_arquivo)
            resultado.documentos.append(nome_arquivo)
        except Exception as exc:  # noqa: BLE001
            resultado.ok = False
            resultado.erros.append(f"{nome_arquivo}: {exc}")

    resultado.pendentes = pendentes
    resultado.alertas = alertas_ipm(dados)
    if _sim(dados.get("sem_carater_disciplinar")):
        resultado.alertas.append(
            "Sindicância sem caráter disciplinar (art. 324, §2º, do MAPPA): a notificação do sindicado "
            "para defesa prévia foi dispensada e a peça correspondente não foi gerada."
        )
    prazo = prazo_conclusao(_data(dados, "data_recibo_portaria"), _sim(dados.get("prorrogado")))
    if prazo:
        resultado.alertas.append(
            f"Prazo de conclusão: {prazo.strftime('%d/%m/%Y')} "
            f"({PRAZO_DIAS}{' + ' + str(PRAZO_PRORROGACAO) if _sim(dados.get('prorrogado')) else ''} "
            f"dias corridos do recibo da portaria - arts. 273/274 do MAPPA)."
        )
    return resultado
